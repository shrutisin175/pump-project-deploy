import os
import io
from datetime import datetime
from decimal import Decimal
from django.http import HttpResponse
from django.conf import settings
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import subprocess
import random

def generate_reference_number(customer_info=None):
    """
    Generate reference number with GST + last 2 digits sequential
    """
    if customer_info and customer_info.get('gst_number'):
        gst = customer_info.get('gst_number')
        # Extract last 2 digits from GST number
        last_two = gst[-2:] if len(gst) >= 2 else '00'
        # Generate unique sequential number using timestamp + random
        timestamp_part = str(int(datetime.now().timestamp()))[-6:]
        random_part = str(random.randint(100, 999))
        return f"{gst}{last_two}{timestamp_part}{random_part}"
    else:
        # Fallback: generate random reference with timestamp
        timestamp_part = datetime.now().strftime('%y%m%d%H%M')
        random_part = str(random.randint(1000, 9999))
        return f"INV{timestamp_part}{random_part}"

def create_comprehensive_template():
    """
    Create the comprehensive pump spares quotation template
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('OFFER FOR SUPPLY OF PUMP SPARES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add reference and date
    doc.add_paragraph('Ref No: [Ref. No.]')
    doc.add_paragraph('Date: [DD/MM/YYYY]')
    
    # Add customer details
    doc.add_paragraph('To:')
    doc.add_paragraph('[Customer Name & Address]')
    doc.add_paragraph('')
    doc.add_paragraph('_' * 40)
    
    # Section 1: Offer Details
    doc.add_heading('1. OFFER DETAILS', level=2)
    doc.add_paragraph('We are pleased to submit our offer for the supply of the following pump spares as per your requirement:')
    
    # Create table for spare details
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sr. No.'
    hdr_cells[1].text = 'Description of Spare'
    hdr_cells[2].text = 'Qty'
    hdr_cells[3].text = 'Unit Price (INR)'
    hdr_cells[4].text = 'Total Price (INR)'
    
    # Add data row placeholder
    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = '[Spare Name & Details]'
    row_cells[2].text = '1'
    row_cells[3].text = '[XXXX]'
    row_cells[4].text = '[XXXX]'
    
    doc.add_paragraph('')
    doc.add_paragraph('Total Basic Price: ₹[XXXX]')
    doc.add_paragraph('GST: Extra as applicable')
    doc.add_paragraph('Packing & Forwarding: Extra at actuals')
    doc.add_paragraph('Freight & Insurance: Extra at actuals or on customer\'s account')
    doc.add_paragraph('Delivery Terms: Ex-works')
    doc.add_paragraph('_' * 40)
    
    # Section 2: Commercial Terms
    doc.add_heading('2. COMMERCIAL TERMS', level=2)
    doc.add_paragraph('1. Validity of Offer: Valid for 30 (thirty) days from the date of this quotation.')
    doc.add_paragraph('2. Price Basis: Prices are based on current raw material rates and are subject to change without prior notice if the order is not placed within the validity period.')
    doc.add_paragraph('3. Taxes & Duties: All taxes, duties, levies (present or future) applicable shall be borne by the buyer. Any increase after the order date will be on the buyer\'s account.')
    doc.add_paragraph('4. Packing & Forwarding: Will be charged extra at actuals.')
    doc.add_paragraph('5. Freight & Insurance:')
    doc.add_paragraph('   o Freight and transit insurance to the customer\'s site will be borne by the buyer unless otherwise mentioned in writing.')
    doc.add_paragraph('   o Risk in goods passes to the buyer upon delivery to the carrier at our works, even if freight is prepaid.')
    doc.add_paragraph('6. Delivery Schedule:')
    doc.add_paragraph('   o Expected delivery: [X weeks] from the date of receipt of firm Purchase Order and advance payment.')
    doc.add_paragraph('   o Delivery schedule is indicative and subject to availability of raw materials, transport, and labor conditions.')
    doc.add_paragraph('   o Delay caused due to reasons beyond our control (Force Majeure) shall not attract any penalty or cancellation.')
    doc.add_paragraph('7. Material Acceptance:')
    doc.add_paragraph('   o The material supplied shall be deemed to have been accepted by the buyer as free from any defect in material, design, manufacturing, or workmanship unless written notice of any such defect is received by us within seven (7) days from the date of delivery at the buyer\'s premises.')
    doc.add_paragraph('   o Proof of Delivery: The delivery report of the courier/transporter or any receipt indicating gate entry to the buyer\'s premises shall be considered conclusive proof of delivery.')
    doc.add_paragraph('   o After this period or once the material is installed/used, no claims for rejection, replacement, or rectification will be entertained.')
    doc.add_paragraph('   o Any refusal to accept delivery shall not relieve the buyer of their payment obligation, and the goods will be considered delivered for invoicing purposes.')
    doc.add_paragraph('_' * 40)
    
    # Section 3: Payment Terms
    doc.add_heading('3. PAYMENT TERMS', level=2)
    doc.add_paragraph('● Advance: 50% of order value along with Purchase Order.')
    doc.add_paragraph('● Balance: Before dispatch / Against proforma invoice.')
    doc.add_paragraph('● Payment Mode: NEFT/RTGS only.')
    doc.add_paragraph('● Delayed Payment: Any delayed payment beyond agreed terms will attract interest @ 18% per annum from the due date till realization.')
    doc.add_paragraph('● Ownership: Goods remain our property until full and final payment is received.')
    doc.add_paragraph('_' * 40)
    
    # Section 4: Warranty Terms
    doc.add_heading('4. WARRANTY TERMS', level=2)
    doc.add_paragraph('● Warranty is applicable for 6 months from supply date or 3 months from commissioning, whichever occurs earlier.')
    doc.add_paragraph('● Warranty covers manufacturing defects only.')
    doc.add_paragraph('● Warranty excludes:')
    doc.add_paragraph('   o Normal wear & tear')
    doc.add_paragraph('   o Misuse, mishandling, improper storage, faulty installation')
    doc.add_paragraph('   o Operation outside design parameters')
    doc.add_paragraph('   o Use of non-genuine spare parts with our supplied spares')
    doc.add_paragraph('● Our liability under warranty is limited to repair/replacement of defective part only.')
    doc.add_paragraph('● Warranty claim does not cover labor cost, travel cost, removal or reinstallation charges, unless agreed in writing.')
    doc.add_paragraph('● Buyer shall provide sufficient evidence (photographs, reports, defective part return) for warranty claim evaluation.')
    doc.add_paragraph('_' * 40)
    
    # Section 5: Cancellation Policy
    doc.add_heading('5. CANCELLATION POLICY', level=2)
    doc.add_paragraph('● Orders once accepted cannot be cancelled without written consent from us.')
    doc.add_paragraph('● For made-to-order or customized spares, 100% cancellation charges will apply once production begins.')
    doc.add_paragraph('_' * 40)
    
    # Section 6: Limitation of Liability
    doc.add_heading('6. LIMITATION OF LIABILITY', level=2)
    doc.add_paragraph('● Our maximum liability shall not exceed the invoice value of defective goods supplied under any circumstances.')
    doc.add_paragraph('● We are not liable for any indirect, incidental, or consequential losses, including but not limited to:')
    doc.add_paragraph('   o Loss of production')
    doc.add_paragraph('   o Loss of profit')
    doc.add_paragraph('   o Delay in project')
    doc.add_paragraph('   o Third-party claims')
    doc.add_paragraph('_' * 40)
    
    # Section 7: Inspection & Acceptance
    doc.add_heading('7. INSPECTION & ACCEPTANCE', level=2)
    doc.add_paragraph('● Pre-dispatch inspection (PDI), if required, should be communicated in advance and will be on buyer\'s account.')
    doc.add_paragraph('● Buyer must inspect goods immediately upon receipt. Any shortage or visible damage must be notified within 48 hours of delivery.')
    doc.add_paragraph('_' * 40)
    
    # Section 8: Force Majeure
    doc.add_heading('8. FORCE MAJEURE', level=2)
    doc.add_paragraph('We shall not be held responsible for delay or failure in fulfilling obligations caused by circumstances beyond our control, including but not limited to strikes, lockouts, transport delays, accidents, natural calamities, fire, floods, acts of God, war, or government restrictions.')
    doc.add_paragraph('_' * 40)
    
    # Section 9: Dispute Resolution & Jurisdiction
    doc.add_heading('9. DISPUTE RESOLUTION & JURISDICTION', level=2)
    doc.add_paragraph('● Any dispute arising out of this contract shall be subject to the exclusive jurisdiction of the courts at [Your City] only.')
    doc.add_paragraph('● Arbitration, if invoked, shall be under the Arbitration & Conciliation Act, 1996, and the seat of arbitration shall be [Your City].')
    doc.add_paragraph('_' * 40)
    
    # Section 10: Other Conditions
    doc.add_heading('10. OTHER CONDITIONS', level=2)
    doc.add_paragraph('● Technical support (if required) will be chargeable extra unless explicitly mentioned otherwise.')
    doc.add_paragraph('● Storage and preservation of spares after delivery is the responsibility of the buyer.')
    doc.add_paragraph('● Any modification to this offer is valid only if confirmed by us in writing.')
    doc.add_paragraph('● Buyer agrees to all terms and conditions mentioned herein by placing a Purchase Order against this offer.')
    doc.add_paragraph('_' * 40)
    
    # Closing
    doc.add_paragraph('We look forward to your confirmation and valuable order.')
    doc.add_paragraph('')
    doc.add_paragraph('For')
    doc.add_paragraph('Shaft & Seal Pvt ltd')
    doc.add_paragraph('Authorized Signatory')
    doc.add_paragraph('Name: [Authorized Signatory Name]')
    doc.add_paragraph('Designation: [Designation]')
    doc.add_paragraph('Contact Details: [Contact Details]')
    
    return doc

def create_inventory_receipt_response(cart_items, customer_info=None):
    """
    Generate a receipt document for inventory items and return as HTTP response
    """
    try:
        # Create comprehensive template
        doc = create_comprehensive_template()
        
        # Generate reference number
        ref_no = generate_reference_number(customer_info)
        current_date = datetime.now().strftime('%d/%m/%Y')
        
        # Auto-populate customer details
        if customer_info:
            customer_name = customer_info.get('name', '')
            customer_address = customer_info.get('address', '')
            customer_email = customer_info.get('email', '')
            customer_location = customer_info.get('location', '')
            
            # Format customer address properly
            customer_full_address = f"{customer_name}"
            if customer_address:
                customer_full_address += f"\n{customer_address}"
            if customer_email:
                customer_full_address += f"\nEmail: {customer_email}"
        else:
            customer_name = 'Valued Customer'
            customer_address = ''
            customer_full_address = 'Valued Customer'
            customer_location = ''
        
        # Calculate total price
        total_basic_price = Decimal('0')
        
        # Find and populate the table
        for table in doc.tables:
            if len(table.rows) > 1 and 'Sr. No.' in table.rows[0].cells[0].text:
                # Clear existing data rows
                while len(table.rows) > 1:
                    table._element.remove(table.rows[1]._element)
                
                # Add cart items
                for idx, item in enumerate(cart_items, 1):
                    row = table.add_row()
                    
                    row.cells[0].text = str(idx)
                    
                    # Create detailed spare description
                    description = f"{item['partName']} - {item['moc']}"
                    if item.get('drawing'):
                        description += f"\nDrawing: {item['drawing']}"
                    description += f"\nPart No: {item['partNo']}"
                    
                    row.cells[1].text = description
                    row.cells[2].text = str(item['quantity'])
                    
                    unit_price = Decimal(str(item['unitPrice']))
                    row.cells[3].text = f"{unit_price:,.2f}"
                    
                    total_price = unit_price * Decimal(str(item['quantity']))
                    row.cells[4].text = f"{total_price:,.2f}"
                    
                    total_basic_price += total_price
                
                break
        
        # Replace placeholders
        replacements = {
            '[Ref. No.]': ref_no,
            '[DD/MM/YYYY]': current_date,
            '[Customer Name & Address]': customer_full_address,
            '[XXXX]': f"{total_basic_price:,.2f}",
            '[X weeks]': '2-3 weeks',
            '[50%]': '50%',
            '[Your City/State]': customer_location,
            '[Your City]': customer_location,
            '[Authorized Signatory Name]': 'Authorized Signatory',
            '[Designation]': 'Sales Manager',
            '[Contact Details]': 'Email: info@shaftseal.com | Phone: +91-XXXXXXXXXX'
        }
        
        # Replace text in document
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace text in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, replacement in replacements.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_content = buffer.getvalue()
        
        # Try to convert to PDF using alternative method
        try:
            pdf_content = convert_docx_to_pdf_alternative(file_content)
            
            response = HttpResponse(
                pdf_content,
                content_type='application/pdf'
            )
            
            filename = f"Pump_Spares_Quotation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as pdf_error:
            # Fallback to DOCX if PDF conversion fails
            print(f"PDF conversion failed, falling back to DOCX: {pdf_error}")
            
            response = HttpResponse(
                file_content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            filename = f"Pump_Spares_Quotation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
        
    except Exception as e:
        error_response = HttpResponse(
            f'Error generating receipt: {str(e)}',
            content_type='text/plain',
            status=500
        )
        return error_response

def convert_docx_to_pdf_alternative(docx_content):
    """
    Convert DOCX to PDF using alternative method (currently returns DOCX)
    """
    # For now, return the DOCX content as PDF is not available without LibreOffice
    raise Exception("PDF conversion not available - LibreOffice not installed")

def convert_docx_to_pdf_bytes(docx_content):
    """
    Convert DOCX bytes to PDF using LibreOffice
    """
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_content)
            temp_docx_path = temp_docx.name
        
        # Create output directory
        output_dir = tempfile.mkdtemp()
        
        # Convert using LibreOffice
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            temp_docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0:
            # Find the generated PDF file
            pdf_filename = os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
            pdf_path = os.path.join(output_dir, pdf_filename)
            
            if os.path.exists(pdf_path):
                # Read the generated PDF
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up temporary files
                os.unlink(temp_docx_path)
                import shutil
                shutil.rmtree(output_dir)
                
                return pdf_content
            else:
                raise Exception(f"PDF file not found: {pdf_path}")
        else:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
    except subprocess.TimeoutExpired:
        raise Exception("PDF conversion timed out")
    except Exception as e:
        # Clean up temporary files if they exist
        try:
            if 'temp_docx_path' in locals():
                os.unlink(temp_docx_path)
            if 'output_dir' in locals():
                import shutil
                shutil.rmtree(output_dir)
        except:
            pass
        raise Exception(f"PDF conversion error: {str(e)}")