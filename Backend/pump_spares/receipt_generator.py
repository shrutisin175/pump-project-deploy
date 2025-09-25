from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.http import HttpResponse
from django.conf import settings
import os
from datetime import datetime
import io
import hashlib
import random
import subprocess
import tempfile

def generate_reference_number(user_info=None):
    """
    Generate reference number with GST + last 2 digits sequential
    """
    if user_info and user_info.get('gst_number'):
        gst = user_info.get('gst_number')
        # Extract last 2 digits from GST number
        last_two = gst[-2:] if len(gst) >= 2 else '00'
        # Generate sequential number (using timestamp for uniqueness)
        sequential = str(int(datetime.now().timestamp()))[-4:]
        return f"{gst}{last_two}{sequential}"
    else:
        # Fallback: generate random reference
        return f"SS{datetime.now().strftime('%y%m%d')}{random.randint(1000, 9999)}"

def create_basic_template():
    """
    Create the comprehensive pump spares quotation template
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('OFFER FOR SUPPLY OF PUMP SPARES', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add reference and date
    doc.add_paragraph('Ref No: [Ref. No.]')
    doc.add_paragraph('Date: [DD/MM/YYYY]')
    
    # Add customer details
    doc.add_paragraph('To:')
    doc.add_paragraph('[Customer Name & Address]')
    doc.add_paragraph('')
    
    # Add main title again
    title2 = doc.add_heading('OFFER FOR SUPPLY OF PUMP SPARES', level=1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Ref No: [Ref. No.]')
    doc.add_paragraph('Date: [DD/MM/YYYY]')
    doc.add_paragraph('To:')
    doc.add_paragraph('')
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('_' * 40)
    
    # Section 1: Offer Details
    doc.add_heading('1. OFFER DETAILS', level=2)
    doc.add_paragraph('We are pleased to submit our offer for the supply of the following pump spares as per your requirement:')
    
    # Create table for spare details
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sr. No.'
    hdr_cells[1].text = 'Description of Spare'
    hdr_cells[2].text = 'Qty'
    hdr_cells[3].text = 'Unit Price (INR)'
    hdr_cells[4].text = 'Total Price (INR)'
    
    # Add data row
    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = '[Spare Name & Details]'
    row_cells[2].text = '1'
    row_cells[3].text = '[XXXX]'
    row_cells[4].text = '[XXXX]'
    
    doc.add_paragraph('')
    doc.add_paragraph('Total Basic Price: ₹[XXXX]')
    doc.add_paragraph('GST: Extra as applicable')
    doc.add_paragraph('Packing & Forwarding: Extra at actuals')
    doc.add_paragraph('Freight & Insurance: Extra at actuals or on customer\'s account')
    doc.add_paragraph('Delivery Terms: Ex-works')
    doc.add_paragraph('_' * 40)
    
    # Section 2: Commercial Terms
    doc.add_heading('2. COMMERCIAL TERMS', level=2)
    doc.add_paragraph('1. Validity of Offer: Valid for 30 (thirty) days from the date of this quotation.')
    doc.add_paragraph('2. Price Basis: Prices are based on current raw material rates and are subject to change without prior notice if the order is not placed within the validity period.')
    doc.add_paragraph('3. Taxes & Duties: All taxes, duties, levies (present or future) applicable shall be borne by the buyer. Any increase after the order date will be on the buyer\'s account.')
    doc.add_paragraph('4. Packing & Forwarding: Will be charged extra at actuals.')
    doc.add_paragraph('5. Freight & Insurance:')
    doc.add_paragraph('   o Freight and transit insurance to the customer\'s site will be borne by the buyer unless otherwise mentioned in writing.')
    doc.add_paragraph('   o Risk in goods passes to the buyer upon delivery to the carrier at our works, even if freight is prepaid.')
    doc.add_paragraph('6. Delivery Schedule:')
    doc.add_paragraph('   o Expected delivery: [X weeks] from the date of receipt of firm Purchase Order and advance payment.')
    doc.add_paragraph('   o Delivery schedule is indicative and subject to availability of raw materials, transport, and labor conditions.')
    doc.add_paragraph('   o Delay caused due to reasons beyond our control (Force Majeure) shall not attract any penalty or cancellation.')
    doc.add_paragraph('7. Material Acceptance:')
    doc.add_paragraph('   o The material supplied shall be deemed to have been accepted by the buyer as free from any defect in material, design, manufacturing, or workmanship unless written notice of any such defect is received by us within seven (7) days from the date of delivery at the buyer\'s premises.')
    doc.add_paragraph('   o Proof of Delivery: The delivery report of the courier/transporter or any receipt indicating gate entry to the buyer\'s premises shall be considered conclusive proof of delivery.')
    doc.add_paragraph('   o After this period or once the material is installed/used, no claims for rejection, replacement, or rectification will be entertained.')
    doc.add_paragraph('   o Any refusal to accept delivery shall not relieve the buyer of their payment obligation, and the goods will be considered delivered for invoicing purposes.')
    doc.add_paragraph('_' * 40)
    
    # Section 3: Payment Terms
    doc.add_heading('3. PAYMENT TERMS', level=2)
    doc.add_paragraph('● Advance: 50% of order value along with Purchase Order.')
    doc.add_paragraph('● Balance: Before dispatch / Against proforma invoice.')
    doc.add_paragraph('● Payment Mode: NEFT/RTGS only.')
    doc.add_paragraph('● Delayed Payment: Any delayed payment beyond agreed terms will attract interest @ 18% per annum from the due date till realization.')
    doc.add_paragraph('● Ownership: Goods remain our property until full and final payment is received.')
    doc.add_paragraph('_' * 40)
    
    # Section 4: Warranty Terms
    doc.add_heading('4. WARRANTY TERMS', level=2)
    doc.add_paragraph('● Warranty is applicable for 6 months from supply date or 3 months from commissioning, whichever occurs earlier.')
    doc.add_paragraph('● Warranty covers manufacturing defects only.')
    doc.add_paragraph('● Warranty excludes:')
    doc.add_paragraph('   o Normal wear & tear')
    doc.add_paragraph('   o Misuse, mishandling, improper storage, faulty installation')
    doc.add_paragraph('   o Operation outside design parameters')
    doc.add_paragraph('   o Use of non-genuine spare parts with our supplied spares')
    doc.add_paragraph('● Our liability under warranty is limited to repair/replacement of defective part only.')
    doc.add_paragraph('● Warranty claim does not cover labor cost, travel cost, removal or reinstallation charges, unless agreed in writing.')
    doc.add_paragraph('● Buyer shall provide sufficient evidence (photographs, reports, defective part return) for warranty claim evaluation.')
    doc.add_paragraph('_' * 40)
    
    # Section 5: Cancellation Policy
    doc.add_heading('5. CANCELLATION POLICY', level=2)
    doc.add_paragraph('● Orders once accepted cannot be cancelled without written consent from us.')
    doc.add_paragraph('● For made-to-order or customized spares, 100% cancellation charges will apply once production begins.')
    doc.add_paragraph('_' * 40)
    
    # Section 6: Limitation of Liability
    doc.add_heading('6. LIMITATION OF LIABILITY', level=2)
    doc.add_paragraph('● Our maximum liability shall not exceed the invoice value of defective goods supplied under any circumstances.')
    doc.add_paragraph('● We are not liable for any indirect, incidental, or consequential losses, including but not limited to:')
    doc.add_paragraph('   o Loss of production')
    doc.add_paragraph('   o Loss of profit')
    doc.add_paragraph('   o Delay in project')
    doc.add_paragraph('   o Third-party claims')
    doc.add_paragraph('_' * 40)
    
    # Section 7: Inspection & Acceptance
    doc.add_heading('7. INSPECTION & ACCEPTANCE', level=2)
    doc.add_paragraph('● Pre-dispatch inspection (PDI), if required, should be communicated in advance and will be on buyer\'s account.')
    doc.add_paragraph('● Buyer must inspect goods immediately upon receipt. Any shortage or visible damage must be notified within 48 hours of delivery.')
    doc.add_paragraph('_' * 40)
    
    # Section 8: Force Majeure
    doc.add_heading('8. FORCE MAJEURE', level=2)
    doc.add_paragraph('We shall not be held responsible for delay or failure in fulfilling obligations caused by circumstances beyond our control, including but not limited to strikes, lockouts, transport delays, accidents, natural calamities, fire, floods, acts of God, war, or government restrictions.')
    doc.add_paragraph('_' * 40)
    
    # Section 9: Dispute Resolution & Jurisdiction
    doc.add_heading('9. DISPUTE RESOLUTION & JURISDICTION', level=2)
    doc.add_paragraph('● Any dispute arising out of this contract shall be subject to the exclusive jurisdiction of the courts at [Your City] only.')
    doc.add_paragraph('● Arbitration, if invoked, shall be under the Arbitration & Conciliation Act, 1996, and the seat of arbitration shall be [Your City].')
    doc.add_paragraph('_' * 40)
    
    # Section 10: Other Conditions
    doc.add_heading('10. OTHER CONDITIONS', level=2)
    doc.add_paragraph('● Technical support (if required) will be chargeable extra unless explicitly mentioned otherwise.')
    doc.add_paragraph('● Storage and preservation of spares after delivery is the responsibility of the buyer.')
    doc.add_paragraph('● Any modification to this offer is valid only if confirmed by us in writing.')
    doc.add_paragraph('● Buyer agrees to all terms and conditions mentioned herein by placing a Purchase Order against this offer.')
    doc.add_paragraph('_' * 40)
    
    # Closing
    doc.add_paragraph('We look forward to your confirmation and valuable order.')
    doc.add_paragraph('')
    doc.add_paragraph('For')
    doc.add_paragraph('Shaft & Seal Pvt ltd')
    doc.add_paragraph('Authorized Signatory')
    doc.add_paragraph('Name: [Authorized Signatory Name]')
    doc.add_paragraph('Designation: [Designation]')
    doc.add_paragraph('Contact Details: [Contact Details]')
    
    return doc

def generate_receipt(material_data, user_info=None):
    """
    Generate a receipt document with filled pump spares details
    """
    template_path = os.path.join(settings.BASE_DIR, 'Receipt.docx')
    
    if not os.path.exists(template_path):
        # Create a basic template if it doesn't exist
        doc = create_basic_template()
    else:
        doc = Document(template_path)
    
    current_date = datetime.now().strftime("%d/%m/%Y")
    
    # Generate reference number with GST + last 2 digits sequential
    ref_no = generate_reference_number(user_info)
    
    # Auto-populate customer details from login
    if user_info:
        customer_name = user_info.get('full_name') or user_info.get('company_name') or ''
        customer_address = user_info.get('company_address') or ''
        customer_email = user_info.get('official_email') or user_info.get('email') or ''
        customer_location = user_info.get('location') or ''
        
        # Format customer address properly
        customer_full_address = f"{customer_name}"
        if customer_address:
            customer_full_address += f"\n{customer_address}"
        if customer_email:
            customer_full_address += f"\nEmail: {customer_email}"
    else:
        customer_name = ''
        customer_address = ''
        customer_full_address = ''
        customer_location = ''
    
    unit_price = float(material_data.get('unit_price', 0))
    quantity = 1 
    total_price = unit_price * quantity
    
    # Format prices correctly
    formatted_unit_price = f"{unit_price:,.2f}"
    formatted_total_price = f"{total_price:,.2f}"
    
    # Create detailed spare description
    spare_description = f"{material_data.get('part_name_value', '')} - {material_data.get('moc', '')}"
    if material_data.get('pump_make_name') and material_data.get('pump_model_name') and material_data.get('pump_size_value'):
        spare_description += f"\nPump: {material_data.get('pump_make_name')} {material_data.get('pump_model_name')} {material_data.get('pump_size_value')}"
    if material_data.get('part_number_value'):
        spare_description += f"\nPart No: {material_data.get('part_number_value')}"
    
    replacements = {
        '[Ref. No.]': ref_no,
        '[DD/MM/YYYY]': current_date,
        '[Customer Name & Address]': customer_full_address,
        '[Spare Name & Details]': spare_description,
        'X': str(quantity),
        'XXXX': formatted_unit_price,
        '[XXXX]': formatted_total_price,
        '[Location]': customer_location,
        '[X weeks]': '2-3 weeks',
        '[50%]': '50%',
        '[Your City/State]': customer_location,
        '[Your City]': customer_location,
        '[Authorized Signatory Name]': 'Authorized Signatory',
        '[Designation]': 'Sales Manager',
        '[Contact Details]': 'Email: info@shaftseal.com | Phone: +91-XXXXXXXXXX'
    }
    
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, replacement in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, replacement)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def convert_docx_to_pdf(docx_buffer):
    """
    Convert DOCX buffer to PDF using LibreOffice
    """
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_buffer.getvalue())
            temp_docx_path = temp_docx.name
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name
        
        # Convert using LibreOffice
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(temp_pdf_path),
            temp_docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            # Read the generated PDF
            with open(temp_pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Clean up temporary files
            os.unlink(temp_docx_path)
            os.unlink(temp_pdf_path)
            
            return pdf_content
        else:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
            
    except Exception as e:
        # Clean up temporary files if they exist
        try:
            if 'temp_docx_path' in locals():
                os.unlink(temp_docx_path)
            if 'temp_pdf_path' in locals():
                os.unlink(temp_pdf_path)
        except:
            pass
        raise Exception(f"PDF conversion error: {str(e)}")

def create_receipt_response(material_data, user_info=None):
    """
    Create HTTP response with the generated receipt document as PDF
    """
    try:
        docx_buffer = generate_receipt(material_data, user_info)
        
        # Try to convert to PDF
        try:
            pdf_content = convert_docx_to_pdf(docx_buffer)
            
            response = HttpResponse(
                pdf_content,
                content_type='application/pdf'
            )
            
            filename = f"Pump_Spares_Quotation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
            
        except Exception as pdf_error:
            # Fallback to DOCX if PDF conversion fails
            print(f"PDF conversion failed, falling back to DOCX: {pdf_error}")
            
            response = HttpResponse(
                docx_buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            filename = f"Pump_Spares_Quotation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            return response
        
    except Exception as e:
        raise Exception(f"Error generating receipt: {str(e)}")
